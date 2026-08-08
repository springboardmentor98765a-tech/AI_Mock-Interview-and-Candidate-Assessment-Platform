"""
resume_parser.py
==================
Module 2 - Resume Parsing & Analysis.

Extracts plain text from an uploaded resume (PDF or DOCX) and derives:
    - Technology / skill detection   -> extract_skills() / extract_skills_by_category()
    - Experience parsing             -> extract_total_experience_years(), extract_work_experience()
    - Education analysis             -> extract_education()
    - Resume summary generation      -> generate_resume_summary()

The extracted skills also feed Module 3 (AI interview generation), so
questions can be based on what's actually on the candidate's resume
instead of a manually-typed domain.
"""

import io
import re
from typing import Optional

from fastapi import HTTPException, status

# ---------------------------------------------------------------------------
# Technology / skill detection
# ---------------------------------------------------------------------------
# Grouped by category so the frontend can render "Languages", "Databases",
# etc. separately instead of one flat list. Matching is case-insensitive
# with word-boundary checks so "Java" doesn't false-match inside
# "JavaScript".
SKILL_CATEGORIES: dict[str, list[str]] = {
    "languages": [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust",
        "Kotlin", "Swift", "PHP", "Ruby",
    ],
    "databases": [
        "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite",
    ],
    "frameworks": [
        "React", "Angular", "Vue", "Node.js", "Express", "Django", "Flask",
        "FastAPI", "Spring Boot", ".NET", "Redux", "Bootstrap", "Tailwind",
    ],
    "cloud_devops": [
        "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "CI/CD",
        "Git", "GitHub", "Jenkins",
    ],
    "data_ml": [
        "Machine Learning", "Deep Learning", "Data Analysis", "Data Science",
        "Pandas", "NumPy", "TensorFlow", "PyTorch", "Scikit-learn",
    ],
    "web_core": [
        "REST API", "GraphQL", "Microservices", "HTML", "CSS",
    ],
    "soft_skills": [
        "Agile", "Scrum", "Jira",
        "Communication", "Leadership", "Problem Solving", "Teamwork",
        "Project Management", "Time Management",
    ],
}

# Flat list kept for backwards compatibility (existing skill-based
# interview generation in ai_question_generator.py just needs a flat list).
KNOWN_SKILLS = [skill for group in SKILL_CATEGORIES.values() for skill in group]


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF parsing library not installed on the server.",
        )

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(text_parts)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read that PDF. Please upload a valid, non-encrypted PDF resume.",
        )


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX parsing library not installed on the server.",
        )

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of any tables (resumes often use them for layout)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read that DOCX file. Please upload a valid Word resume.",
        )


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    lower_name = (filename or "").lower()

    if lower_name.endswith(".pdf"):
        text = _extract_pdf_text(file_bytes)
    elif lower_name.endswith(".docx"):
        text = _extract_docx_text(file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type. Please upload a .pdf or .docx resume.",
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Couldn't find any readable text in that file (it may be a scanned image resume).",
        )

    return text


def extract_skills(resume_text: str) -> list[str]:
    """Matches resume text against KNOWN_SKILLS, word-boundary safe."""
    found = []
    for skill in KNOWN_SKILLS:
        pattern = r"(?<![A-Za-z0-9])" + re.escape(skill) + r"(?![A-Za-z0-9])"
        if re.search(pattern, resume_text, flags=re.IGNORECASE):
            found.append(skill)
    return found


def extract_skills_by_category(resume_text: str) -> dict[str, list[str]]:
    """Same matching as extract_skills(), but grouped by category, e.g.
    {"languages": ["Python", "Java"], "databases": ["PostgreSQL"], ...}.
    Empty categories are omitted."""
    result: dict[str, list[str]] = {}
    for category, skills in SKILL_CATEGORIES.items():
        found = []
        for skill in skills:
            pattern = r"(?<![A-Za-z0-9])" + re.escape(skill) + r"(?![A-Za-z0-9])"
            if re.search(pattern, resume_text, flags=re.IGNORECASE):
                found.append(skill)
        if found:
            result[category] = found
    return result


# ---------------------------------------------------------------------------
# Section splitting - shared helper for experience / education parsing
# ---------------------------------------------------------------------------
_SECTION_HEADERS = {
    "summary": ["summary", "profile", "objective", "professional summary"],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment history", "career history", "work history",
    ],
    "education": [
        "education", "academic background", "academic qualifications",
        "educational qualifications", "academics",
    ],
    "skills": ["skills", "technical skills", "core competencies", "key skills"],
    "projects": ["projects", "academic projects", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}
_ALL_HEADER_STRINGS = {
    alias: canonical
    for canonical, aliases in _SECTION_HEADERS.items()
    for alias in aliases
}


def _split_sections(resume_text: str) -> dict[str, str]:
    """Splits resume text into sections keyed by canonical section name
    (e.g. "experience", "education"), based on short lines that look like
    section headers. Text before the first recognized header is returned
    under the key "header" (name, contact info, etc.)."""
    sections: dict[str, list[str]] = {}
    current = "header"
    for raw_line in resume_text.split("\n"):
        stripped = raw_line.strip()
        lookup = re.sub(r"[:\-–—\s]+$", "", stripped).strip().lower()
        if lookup in _ALL_HEADER_STRINGS and len(stripped) <= 40:
            current = _ALL_HEADER_STRINGS[lookup]
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(raw_line)
    return {name: "\n".join(lines).strip() for name, lines in sections.items()}


# ---------------------------------------------------------------------------
# Experience parsing
# ---------------------------------------------------------------------------
_TOTAL_EXPERIENCE_PATTERNS = [
    r"(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of)?\s*(?:relevant|professional|total|work)?\s*experience",
    r"experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*years?",
]

_DATE_RANGE_PATTERN = re.compile(
    r"(?P<start>(?:[A-Za-z]{3,9}\.?\s+)?(?:19|20)\d{2})\s*(?:-|–|—|to)\s*"
    r"(?P<end>(?:[A-Za-z]{3,9}\.?\s+)?(?:19|20)\d{2}|present|current|ongoing|now)",
    flags=re.IGNORECASE,
)


def extract_total_experience_years(resume_text: str) -> Optional[float]:
    """Looks for explicit phrases like '4 years of experience' or
    'Experience: 3+ years' anywhere in the resume. Returns the largest
    figure found (resumes sometimes restate it in a summary and a
    skills line), or None if no such phrase is present."""
    best: Optional[float] = None
    for pattern in _TOTAL_EXPERIENCE_PATTERNS:
        for match in re.finditer(pattern, resume_text, flags=re.IGNORECASE):
            try:
                value = float(match.group(1))
            except (ValueError, IndexError):
                continue
            if best is None or value > best:
                best = value
    return best


def _guess_role_and_company(line: str) -> tuple[str, str]:
    """Given the header line of a job entry (commonly "Title, Company" or
    "Title at Company" or "Title - Company"), splits it into a best-guess
    (title, company) pair. Falls back to (line, "") if no separator found."""
    for sep_pattern in [r"\s+at\s+", r"\s*[|,]\s*", r"\s*-\s*", r"\s*–\s*"]:
        parts = re.split(sep_pattern, line, maxsplit=1, flags=re.IGNORECASE)
        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
            return parts[0].strip(), parts[1].strip()
    return line.strip(), ""


def extract_work_experience(resume_text: str) -> list[dict]:
    """Heuristically extracts individual job entries from the resume's
    Experience section. For each line containing a recognizable date
    range (e.g. "Jan 2021 - Present", "2019 - 2022"), the nearby text is
    used to guess the role/company and duration.

    Returns a list of dicts: [{"title", "company", "duration"}, ...]
    Best-effort only - resumes are not standardized, so this won't catch
    every format, but it works well for the common "role, company (dates)"
    or "role – company \\n dates" layouts.
    """
    sections = _split_sections(resume_text)
    text = sections.get("experience") or resume_text
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    entries = []
    for i, line in enumerate(lines):
        match = _DATE_RANGE_PATTERN.search(line)
        if not match:
            continue

        duration = f"{match.group('start')} - {match.group('end')}"
        # Text on this same line before the date range is often the
        # role/company; if the date range is the whole line, the role is
        # usually on the previous non-empty line instead.
        before_date = line[: match.start()].strip(" -–—|,\t")
        header_line = before_date if before_date else (lines[i - 1] if i > 0 else "")

        if not header_line:
            continue

        title, company = _guess_role_and_company(header_line)
        entries.append({"title": title, "company": company, "duration": duration})

    # De-duplicate consecutive identical entries (can happen if a role
    # and its date both matched on adjacent lines).
    deduped = []
    for entry in entries:
        if not deduped or deduped[-1] != entry:
            deduped.append(entry)
    return deduped


# ---------------------------------------------------------------------------
# Education analysis
# ---------------------------------------------------------------------------
_DEGREE_PATTERN = re.compile(
    r"\b("
    r"B\.?\s?Tech|B\.?\s?E\.?|M\.?\s?Tech|MBA|MCA|BCA|B\.?\s?Sc\.?|M\.?\s?Sc\.?|"
    r"B\.?\s?Com|M\.?\s?Com|Ph\.?\s?D|"
    r"Bachelor(?:'s)?\s+of\s+[A-Za-z][A-Za-z .]*|"
    r"Master(?:'s)?\s+of\s+[A-Za-z][A-Za-z .]*|"
    r"Associate(?:'s)?\s+Degree(?:\s+in\s+[A-Za-z][A-Za-z .]*)?|"
    r"Diploma(?:\s+in\s+[A-Za-z][A-Za-z .]*)?|"
    r"High\s+School(?:\s+Diploma)?"
    r")\b",
    flags=re.IGNORECASE,
)
_YEAR_PATTERN = re.compile(r"\b((?:19|20)\d{2})\b")


def extract_education(resume_text: str) -> list[dict]:
    """Heuristically extracts education entries: degree, institution
    (best guess), and graduation year, primarily by scanning the
    resume's Education section (falling back to the whole document if no
    such section header was found).

    Returns a list of dicts: [{"degree", "institution", "year"}, ...]
    """
    sections = _split_sections(resume_text)
    text = sections.get("education") or resume_text
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    entries = []
    for i, line in enumerate(lines):
        degree_match = _DEGREE_PATTERN.search(line)
        if not degree_match:
            continue

        degree = re.sub(r"\s+", " ", degree_match.group(1)).strip()

        # Institution: rest of the same line (minus the degree text), or
        # the next non-empty line if nothing else is on this one.
        remainder = (line[: degree_match.start()] + line[degree_match.end():]).strip(" ,-–—|\t")
        # Common pattern: "B.Tech in Computer Science, ABC University" -
        # strip the leading "in <field of study>," clause so what's left
        # is just the institution name.
        remainder = re.sub(r"^in\s+[A-Za-z][A-Za-z &]*,\s*", "", remainder, flags=re.IGNORECASE)
        institution = remainder
        if not institution and i + 1 < len(lines):
            institution = lines[i + 1]
        institution = re.sub(r"\b((?:19|20)\d{2})\b", "", institution).strip(" ,-–—|\t")

        # Graduation year: look on this line first, then the next line.
        year_match = _YEAR_PATTERN.search(line) or (
            _YEAR_PATTERN.search(lines[i + 1]) if i + 1 < len(lines) else None
        )
        year = year_match.group(1) if year_match else None

        entries.append({"degree": degree, "institution": institution or None, "year": year})

    # De-duplicate identical degree+institution pairs.
    deduped = []
    seen = set()
    for entry in entries:
        key = (entry["degree"].lower(), (entry["institution"] or "").lower())
        if key not in seen:
            seen.add(key)
            deduped.append(entry)
    return deduped


# ---------------------------------------------------------------------------
# Resume summary generation
# ---------------------------------------------------------------------------
def _generate_summary_with_gemini(resume_text: str) -> Optional[str]:
    from app.config import settings

    if not settings.GEMINI_API_KEY:
        return None

    try:
        import google.generativeai as genai

        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel(settings.GEMINI_MODEL)
        prompt = (
            "Write a concise, professional 2-3 sentence summary of the "
            "candidate described by the resume text below. Third person, "
            "factual, no headings, no markdown - plain prose only.\n\n"
            f"Resume text:\n{resume_text[:6000]}"
        )
        response = model.generate_content(prompt)
        summary = (response.text or "").strip()
        return summary or None
    except Exception:
        return None


def _generate_summary_locally(
    skills: list[str],
    experience_years: Optional[float],
    experience: list[dict],
    education: list[dict],
) -> str:
    parts = []

    if experience_years:
        years_text = f"{experience_years:g} years of experience"
    elif experience:
        years_text = f"experience across {len(experience)} role{'s' if len(experience) != 1 else ''}"
    else:
        years_text = None

    if years_text and skills:
        parts.append(
            f"Candidate with {years_text}, skilled in {', '.join(skills[:6])}."
        )
    elif years_text:
        parts.append(f"Candidate with {years_text}.")
    elif skills:
        parts.append(f"Candidate skilled in {', '.join(skills[:6])}.")
    else:
        parts.append("Candidate profile summarized from the uploaded resume.")

    if experience:
        latest = experience[0]
        role_bit = latest["title"] + (f" at {latest['company']}" if latest.get("company") else "")
        parts.append(f"Most recent role: {role_bit}.")

    if education:
        edu = education[0]
        edu_bit = edu["degree"] + (f" from {edu['institution']}" if edu.get("institution") else "")
        parts.append(f"Education: {edu_bit}.")

    return " ".join(parts)


def generate_resume_summary(
    resume_text: str,
    skills: list[str],
    experience_years: Optional[float] = None,
    experience: Optional[list[dict]] = None,
    education: Optional[list[dict]] = None,
) -> str:
    """Generates a short human-readable summary of the resume. Tries
    Gemini first for a natural-language summary; falls back to a
    template built from the already-extracted structured fields so this
    never hard-fails even with no API key configured."""
    summary = _generate_summary_with_gemini(resume_text)
    if summary:
        return summary
    return _generate_summary_locally(skills, experience_years, experience or [], education or [])


def compute_resume_score(
    skills: list[str],
    experience_years: Optional[float] = None,
    experience: Optional[list[dict]] = None,
    education: Optional[list[dict]] = None,
    summary: Optional[str] = None,
) -> float:
    """
    Real, deterministic resume-completeness score (0-100), computed from
    what was actually extracted from the uploaded file - not a fixed
    placeholder. Weighted across:
        - breadth of detected skills        (up to 40 pts)
        - work experience detected          (up to 25 pts)
        - years of experience detected      (up to 15 pts)
        - education detected                (up to 15 pts)
        - a generated summary exists        (up to 5 pts)
    """
    experience = experience or []
    education = education or []

    skills_score = min(40, len(skills or []) * 4)
    experience_entries_score = min(25, len(experience) * 12)
    years_score = min(15, (experience_years or 0) * 3)
    education_score = 15 if education else 0
    summary_score = 5 if summary else 0

    total = skills_score + experience_entries_score + years_score + education_score + summary_score
    return round(max(0, min(100, total)), 1)
